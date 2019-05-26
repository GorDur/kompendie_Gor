# Writhing a word file bibliotek
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
# Reading an excel file using Python bibliotek
import xlrd   
# importerar ch definerar tid
import datetime
dt = datetime.date.today()
# Give the location of the file 
loc = ("uppgift_3//Bedömning_prog.xlsx")  # kan ändras beroende på positionen 
# För att välja vilken ska tas fram
wb = xlrd.open_workbook(loc) 
sheet = wb.sheet_by_index(0)  
sheet1 = wb.sheet_by_index(1)  
sheet2 = wb.sheet_by_index(2)  
sheet3 = wb.sheet_by_index(3)  
# variabler för klass,ämne och lärare
klass = sheet3.cell_value(2, 1)
larare = sheet3.cell_value(0, 1)
amne = sheet3.cell_value(1, 1)
# documenten funktion till variabel
document = Document()  
f = []  # array för alla elev numer
# elevens namn
elev = input("Ange elevens namn: ")
# for loop för att ta alla elever 
elever = []
for i in range(1, 30):     
    elever.append(sheet.cell_value(i,0))
# rubrik
document.add_heading(amne, 0)
# skriver ut klass och elevens namn
document.add_heading('Klass: '+ klass + "     " +  'Elev: '+ elev) 
# headers
document.add_heading('Centralt innehåll', level = 1)
# skriver i header
scetion = document.sections[0]
header = scetion.header
paragraph = header.paragraphs[0]
paragraph.text = larare + "          " + str(dt)
# elevens position från listan  med alla elever
ID = elever.index(elev) + 1
# information om ämnesbedömning
krav = []    
for i in range(1, 9):  # tar alla värde från denna rad (kraven)
    krav.append(sheet.cell_value(0, i)) # Lägger in i arrayen
# for loop för att gå igenom kraven och sedan skriva ut som List Bullet 
p = []  
for i in range(0, len(krav)):   
    a = document.add_paragraph(style='List Bullet')
    p.append(a.add_run(krav[i]))    # och lägga i en array
# tittar om eleven har genomfört denna delen 
for i in range(1, 9):   
    if sheet.cell_value(ID, i).startswith("Genomfört"):   # Om delen är genomfört då skrevs det fet stilad
        p[i-1].bold = True
# kunskapskrav
document.add_heading('Kunskapskrav')
# skapar tabell
table = document.add_table(rows=11, cols=4)
hdr_cells = table.rows[0].cells # säter variabel för första raden i tabellet
ftr_cells = table.columns[0].cells  # säter variabel för första kolumnen i tabellet
ftr1_cells = table.columns[1].cells
ftr2_cells = table.columns[2].cells
ftr3_cells = table.columns[3].cells
hdr_cells[0].text = 'Rubrik'    # sätter konstanta värde på några ställe
hdr_cells[1].text = 'E'
hdr_cells[2].text = 'C'
hdr_cells[3].text = 'A'
for i in range(0,11):   # vad ska finnas i alla 4 kolumner
    ftr_cells[i].text = sheet2.cell_value(i, 0)
    ftr1_cells[i].text = sheet2.cell_value(i, 1)
    ftr2_cells[i].text = sheet2.cell_value(i, 2)
    ftr3_cells[i].text = sheet2.cell_value(i, 3)
# arrayer för betyg
A = []
C = []
E = []
for i in range(1,11):   # sätter in värde i arrayer
    ID = elever.index(elev) + 1
    if sheet1.cell_value(ID, i).startswith("A"):
        A.append(i)
    elif sheet1.cell_value(ID, i).startswith("C"):
        C.append(i)
    elif sheet1.cell_value(ID, i).startswith("E"):
        E.append(i)
# gör så att rätt kunskap nivå blir färgad
for s in range(1,len(A)+1): # för A nivå
    a = []  # array för att få det som står i tabelen
    a.append(table.cell(A[s-1],3))
    paragraphs = a[0].paragraphs    # paragraph som tar värdet från arrayen
    for paragraph in paragraphs:
        for run in paragraph.runs:  # gör run objekt
            font = run.font 
            font.bold = True    # gör texten fet
            font.highlight_color = 4    # väljer highlight färgen
for s in range (1,len(C)+1):    # för C nivå
    c = []
    c.append(table.cell(C[s-1],2))
    paragraphs = c[0].paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.bold = True
            font.highlight_color = 4    
for s in range (1,len(E)+1):    # för E nivå
    e = []
    e.append(table.cell(E[s-1],1))
    paragraphs = e[0].paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.bold = True
            font.highlight_color = 4     
# skapar documentet om elevens namn finns på listan
if elev in elever:                    
    document.save(elev + ".docx")
else:   # annars skriver ut fel
    print("Fel namn.")